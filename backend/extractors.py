import io
import re
import logging
from typing import Optional

logger = logging.getLogger(__name__)

try:
    import pypdfium2 as pdfium
except Exception:
    pdfium = None

try:
    import docx as _docx_module
except Exception:
    _docx_module = None


DEPT_KEYWORDS: dict = {
    "Computer Technology": [
        "DSA", "C", "C++", "Python", "Java",
        "Data Structures", "Algorithms", "OOP",
        "Operating Systems", "DBMS", "Computer Networks",
    ],
    "Information Technology": [
        "Web Development", "React", "Node.js", "SQL", "Python",
        "APIs", "Cloud", "Docker", "Kubernetes", "JavaScript",
        "HTML", "CSS", "REST", "MongoDB", "AWS",
    ],
}


def extract_text(file_bytes: bytes, mime_type: str) -> str:
    mt = str(mime_type or "").strip().lower()
    if mt == "application/pdf":
        return _extract_pdf(file_bytes)
    if mt in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _extract_docx(file_bytes)
    return ""


def _extract_pdf(file_bytes: bytes) -> str:
    if not pdfium:
        return ""
    try:
        doc = pdfium.PdfDocument(file_bytes)
        parts = []
        for i in range(len(doc)):
            page = doc[i]
            textpage = page.get_textpage()
            parts.append(textpage.get_text_bounded())
        return "\n".join(filter(None, parts))
    except Exception:
        logger.exception("PDF text extraction failed")
        return ""


def _extract_docx(file_bytes: bytes) -> str:
    if not _docx_module:
        return ""
    try:
        doc = _docx_module.Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        logger.exception("DOCX text extraction failed")
        return ""


def _tokenize(text: str) -> list:
    return re.findall(r"[a-zA-Z0-9_.+#-]+", text.lower())


def score_placements_by_query(placements: list, query: str) -> list:
    """
    Score placements by keyword relevance using TF-style scoring.

    Each item in `placements` must have keys: id, extracted_text, company_name.
    Returns items sorted by descending score. If no item scores > 0, returns original order.
    """
    query_tokens = set(_tokenize(query))
    if not query_tokens:
        return placements

    scored = []
    for p in placements:
        corpus = " ".join(filter(None, [
            p.get("extracted_text") or "",
            p.get("company_name") or "",
        ]))
        doc_tokens = _tokenize(corpus)
        if not doc_tokens:
            scored.append((0.0, p))
            continue
        total = len(doc_tokens)
        freq: dict = {}
        for t in doc_tokens:
            freq[t] = freq.get(t, 0) + 1
        score = sum(freq.get(qt, 0) / total for qt in query_tokens)
        scored.append((score, p))

    max_score = max((s for s, _ in scored), default=0.0)
    if max_score == 0.0:
        return placements

    scored.sort(key=lambda x: x[0], reverse=True)
    return [p for s, p in scored if s > 0]
